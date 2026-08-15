
import hashlib
import hmac
import json
import os
import sqlite3
import tempfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import List, Literal, Optional

import fitz
import jwt
import numpy as np
import requests
from docx import Document
from dotenv import load_dotenv
from fastapi import (
    Depends,
    FastAPI,
    File,
    Header,
    HTTPException,
    UploadFile,
)
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, field_validator


# ============================================================
# ENVIRONMENT
# ============================================================

load_dotenv()

ROOT = Path(__file__).resolve().parents[2]

DB_PATH = ROOT / "quizmate.db"
TOPICS_DIR = ROOT / "data" / "topics"

JWT_SECRET = os.getenv(
    "JWT_SECRET",
    "7G5F7aYahj4QHHNK2aALfFcVNGgtNb1M_TnlYIy43Sg",
)

JWT_ALG = "HS256"

FRONTEND_ORIGIN = os.getenv(
    "FRONTEND_ORIGIN",
    "http://localhost:5173",
)

OLLAMA_BASE_URL = os.getenv(
    "OLLAMA_BASE_URL",
    "http://localhost:11434",
).rstrip("/")

OLLAMA_MODEL = os.getenv(
    "OLLAMA_MODEL",
    "llama3.2:latest",
)

OLLAMA_EMBED_MODEL = os.getenv(
    "OLLAMA_EMBED_MODEL",
    "nomic-embed-text:latest",
)


# ============================================================
# FASTAPI
# ============================================================

app = FastAPI(
    title="QuizMate AI API",
    version="1.0.0",
)

origins = [
    o.strip()
    for o in FRONTEND_ORIGIN.split(",")
    if o.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================
# DATABASE
# ============================================================

def get_db():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    return con


def init_db():
    con = get_db()

    con.execute(
        """
        CREATE TABLE IF NOT EXISTS users(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL
        )
        """
    )

    con.execute(
        """
        CREATE TABLE IF NOT EXISTS results(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            topic TEXT NOT NULL,
            score INTEGER NOT NULL,
            total INTEGER NOT NULL,
            difficulty TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
        """
    )

    con.commit()
    con.close()


init_db()


# ============================================================
# PYDANTIC MODELS
# ============================================================

class AuthIn(BaseModel):
    name: Optional[str] = None
    email: str
    password: str


class TokenOut(BaseModel):
    token: str
    user: dict


class QuizRequest(BaseModel):
    source_text: Optional[str] = None
    source_name: Optional[str] = None
    topic: Optional[str] = None

    count: int = Field(
        default=10,
        ge=1,
        le=20,
    )

    question_type: Literal[
        "MCQ",
        "True / False",
        "Short Answer",
        "Mixed",
    ]

    difficulty: Literal[
        "Easy",
        "Medium",
        "Hard",
        "Mixed",
    ]


class QuizQuestion(BaseModel):
    type: Literal[
        "MCQ",
        "True / False",
        "Short Answer",
    ]

    question: str

    options: List[str] = Field(
        default_factory=list
    )

    answer: str

    explanation: str

    keywords: List[str] = Field(
        default_factory=list
    )

    @field_validator("type", mode="before")
    @classmethod
    def normalize_type(cls, value):
        if not isinstance(value, str):
            return value

        value = value.strip()

        normalized = value.lower().replace(" ", "")

        if normalized in {
            "true/false",
            "truefalse",
        }:
            return "True / False"

        if normalized == "mcq":
            return "MCQ"

        if normalized in {
            "shortanswer",
            "short-answer",
        }:
            return "Short Answer"

        return value


class QuizOut(BaseModel):
    questions: List[QuizQuestion]


class SubmitRequest(BaseModel):
    source_name: str
    difficulty: str
    questions: List[QuizQuestion]
    answers: dict


class SummaryRequest(BaseModel):
    text: str = Field(
        min_length=20
    )

    mode: Literal[
        "Quick Summary",
        "Detailed Summary",
        "Exam Revision Notes",
        "Explain Simply",
    ]


# ============================================================
# AUTHENTICATION
# ============================================================

def hash_password(
    password: str,
    salt: Optional[bytes] = None,
):
    salt = salt or os.urandom(16)

    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode(),
        salt,
        120_000,
    )

    return salt.hex() + ":" + digest.hex()


def verify_password(
    password: str,
    stored: str,
):
    try:
        salt_hex, digest_hex = stored.split(":")

        test = hashlib.pbkdf2_hmac(
            "sha256",
            password.encode(),
            bytes.fromhex(salt_hex),
            120_000,
        )

        return hmac.compare_digest(
            test.hex(),
            digest_hex,
        )

    except Exception:
        return False


def make_token(
    user_id: int,
    email: str,
):
    payload = {
        "sub": str(user_id),
        "email": email,
        "exp": datetime.now(timezone.utc)
        + timedelta(hours=12),
    }

    return jwt.encode(
        payload,
        JWT_SECRET,
        algorithm=JWT_ALG,
    )


def current_user(
    authorization: Optional[str] = Header(
        default=None
    ),
):
    if (
        not authorization
        or not authorization.startswith("Bearer ")
    ):
        raise HTTPException(
            status_code=401,
            detail="Please log in.",
        )

    token = authorization.split(
        " ",
        1,
    )[1]

    try:
        payload = jwt.decode(
            token,
            JWT_SECRET,
            algorithms=[JWT_ALG],
        )

        return {
            "id": int(payload["sub"]),
            "email": payload["email"],
        }

    except Exception:
        raise HTTPException(
            status_code=401,
            detail="Session expired. Please log in.",
        )


# ============================================================
# TOPICS
# ============================================================

def load_topics():
    TOPICS_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    return {
        p.stem.replace("_", " ").title():
        p.read_text(encoding="utf-8")
        for p in TOPICS_DIR.glob("*.txt")
    }


def get_topic_source(topic: str):
    data = load_topics()

    if topic in data:
        return topic, data[topic]

    topic_lower = topic.strip().lower()

    for name, text in data.items():
        if name.lower() == topic_lower:
            return name, text

    raise HTTPException(
        status_code=404,
        detail=f"Topic '{topic}' was not found.",
    )


# ============================================================
# OLLAMA
# ============================================================

RAG_CACHE = {}


def ollama_post(
    endpoint: str,
    payload: dict,
    timeout: int = 180,
):
    url = f"{OLLAMA_BASE_URL}{endpoint}"

    try:
        response = requests.post(
            url,
            json=payload,
            timeout=timeout,
        )

        response.raise_for_status()

        return response.json()

    except requests.exceptions.ConnectionError:
        raise HTTPException(
            status_code=503,
            detail=(
                "Ollama is not running. "
                "Start Ollama and make sure the required "
                "models are installed."
            ),
        )

    except requests.exceptions.Timeout:
        raise HTTPException(
            status_code=504,
            detail=(
                "Ollama took too long to respond. "
                "Please try again."
            ),
        )

    except requests.exceptions.HTTPError as e:
        detail = ""

        try:
            detail = response.json().get(
                "error",
                "",
            )
        except Exception:
            pass

        raise HTTPException(
            status_code=502,
            detail=(
                f"Ollama API error: "
                f"{detail or str(e)}"
            ),
        )

    except requests.exceptions.RequestException as e:
        raise HTTPException(
            status_code=502,
            detail=f"Ollama request failed: {str(e)}",
        )


def call_ai(prompt: str):
    print(
        "\n========== OLLAMA DEBUG =========="
    )

    print(
        "OLLAMA URL:",
        OLLAMA_BASE_URL,
    )

    print(
        "LLM MODEL:",
        OLLAMA_MODEL,
    )

    print(
        "EMBED MODEL:",
        OLLAMA_EMBED_MODEL,
    )

    print(
        "PROMPT LENGTH:",
        len(prompt),
    )

    payload = {
        "model": OLLAMA_MODEL,

        "messages": [
            {
                "role": "system",
                "content": (
                    "You are QuizMate AI. "
                    "Follow the user's instructions exactly. "
                    "Return valid JSON when requested. "
                    "Do not return markdown."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],

        "stream": False,

        "format": "json",

        "options": {
            "temperature": 0.2,
        },
    }

    data = ollama_post(
        "/api/chat",
        payload,
        timeout=300,
    )

    result = (
        data
        .get("message", {})
        .get("content", "")
    )

    print(
        "OLLAMA SUCCESS:",
        bool(result),
    )

    print(
        "RESPONSE LENGTH:",
        len(result),
    )

    print(
        "RESPONSE PREVIEW:",
        result[:1000],
    )

    print(
        "==================================\n"
    )

    if not result.strip():
        raise HTTPException(
            status_code=502,
            detail=(
                "Ollama returned an empty response."
            ),
        )

    return result


# ============================================================
# JSON HELPERS
# ============================================================

def clean_json(raw: str):
    if not raw:
        return ""

    raw = raw.strip()

    if raw.startswith("```json"):
        raw = raw[len("```json"):]

    elif raw.startswith("```"):
        raw = raw[len("```"):]

    if raw.endswith("```"):
        raw = raw[:-3]

    return raw.strip()


def parse_ai_json(raw: str):
    cleaned = clean_json(raw)

    if not cleaned:
        raise ValueError(
            "AI returned an empty response."
        )

    try:
        return json.loads(cleaned)

    except json.JSONDecodeError as e:
        print(
            "JSON PARSE ERROR:",
            repr(e),
        )

        print(
            "RAW AI RESPONSE:",
            cleaned[:5000],
        )

        raise ValueError(
            "AI returned invalid JSON."
        )


# ============================================================
# QUESTION NORMALIZATION
# ============================================================

def normalize_question_item(
    item: dict,
):
    """
    Normalize common Ollama mistakes before Pydantic validation.
    """

    if not isinstance(item, dict):
        raise ValueError(
            "Question is not a JSON object."
        )

    item = dict(item)

    # -------------------------
    # Normalize type
    # -------------------------

    raw_type = item.get(
        "type",
        "",
    )

    if isinstance(raw_type, str):
        normalized_type = (
            raw_type
            .strip()
            .lower()
            .replace(" ", "")
        )

        if normalized_type in {
            "true/false",
            "truefalse",
        }:
            item["type"] = "True / False"

        elif normalized_type == "mcq":
            item["type"] = "MCQ"

        elif normalized_type in {
            "shortanswer",
            "short-answer",
        }:
            item["type"] = "Short Answer"

    # -------------------------
    # Options
    # -------------------------

    options = item.get(
        "options",
        None,
    )

    if not isinstance(options, list):
        options = []

    item["options"] = [
        str(option).strip()
        for option in options
    ]

    # True/False always gets options
    if item.get("type") == "True / False":
        item["options"] = [
            "True",
            "False",
        ]

    # Short Answer always gets empty options
    if item.get("type") == "Short Answer":
        item["options"] = []

    # -------------------------
    # Keywords
    # -------------------------

    keywords = item.get(
        "keywords",
        [],
    )

    if not isinstance(keywords, list):
        keywords = []

    item["keywords"] = [
        str(k).strip()
        for k in keywords
        if str(k).strip()
    ]

    # -------------------------
    # Required strings
    # -------------------------

    item["question"] = str(
        item.get(
            "question",
            "",
        )
    ).strip()

    item["answer"] = str(
        item.get(
            "answer",
            "",
        )
    ).strip()

    item["explanation"] = str(
        item.get(
            "explanation",
            "",
        )
    ).strip()

    return item


# ============================================================
# QUESTION VALIDATION
# ============================================================

def validate_questions(
    items,
    requested_count: int,
    selected_type: str,
):
    if not isinstance(items, list):
        raise ValueError(
            "AI did not return a questions list."
        )

    validated = []
    seen = set()

    for raw_item in items:

        try:
            normalized = normalize_question_item(
                raw_item
            )

            # =================================================
            # MIXED MODE
            # =================================================

            if selected_type == "Mixed":

                if normalized.get("type") not in {
                    "MCQ",
                    "True / False",
                    "Short Answer",
                }:
                    print(
                        "Skipping unsupported "
                        "question type:",
                        normalized.get("type"),
                    )
                    continue

            # =================================================
            # STRICT SINGLE TYPE MODE
            # =================================================

            else:

                if normalized.get("type") != selected_type:
                    print(
                        "QUESTION TYPE SKIPPED:",
                        normalized.get("type"),
                        "| EXPECTED:",
                        selected_type,
                    )
                    continue

            question = QuizQuestion(
                **normalized
            )

        except Exception as e:
            print(
                "QUESTION VALIDATION SKIPPED:",
                repr(e),
            )
            continue

        question_key = " ".join(
            question.question
            .lower()
            .split()
        )

        if not question_key:
            continue

        if question_key in seen:
            print(
                "Skipping duplicate question."
            )
            continue

        # ====================================================
        # MCQ
        # ====================================================

        if question.type == "MCQ":

            if len(question.options) != 4:
                print(
                    "Skipping MCQ: "
                    "must have exactly 4 options."
                )
                continue

            normalized_options = [
                option.strip().lower()
                for option in question.options
            ]

            if len(
                set(normalized_options)
            ) != 4:
                print(
                    "Skipping MCQ: "
                    "duplicate options."
                )
                continue

            if not question.answer.strip():
                print(
                    "Skipping MCQ: "
                    "answer is empty."
                )
                continue

            if (
                question.answer
                .strip()
                .lower()
                not in normalized_options
            ):
                print(
                    "Skipping MCQ: "
                    "answer is not one of options."
                )
                continue

        # ====================================================
        # TRUE / FALSE
        # ====================================================

        elif question.type == "True / False":

            if (
                question.answer
                .strip()
                .lower()
                not in {
                    "true",
                    "false",
                }
            ):
                print(
                    "Skipping True/False: "
                    "answer must be True or False."
                )
                continue

            question.options = [
                "True",
                "False",
            ]

        # ====================================================
        # SHORT ANSWER
        # ====================================================

        elif question.type == "Short Answer":

            if not question.answer.strip():
                print(
                    "Skipping Short Answer: "
                    "answer is empty."
                )
                continue

            question.options = []

        # ====================================================
        # UNKNOWN TYPE
        # ====================================================

        else:

            print(
                "Skipping unsupported question type:",
                question.type,
            )

            continue

        seen.add(question_key)

        validated.append(question)

        if len(validated) >= requested_count:
            break

    return validated


# ============================================================
# ROBUST AI QUESTION GENERATION
# ============================================================

def generate_validated_questions(
    prompt: str,
    requested_count: int,
    selected_type: str,
    mode: str,
    max_attempts: int = 3,
):
    """
    Ask Ollama for a quiz and retry when the model returns
    malformed, incomplete, duplicate, or wrong-type questions.

    This keeps the existing validation rules but avoids failing
    immediately when Llama produces only a few valid questions.
    """
    last_valid_count = 0
    last_error = ""

    for attempt in range(1, max_attempts + 1):
        retry_instruction = ""

        if attempt > 1:
            retry_instruction = f"""

IMPORTANT RETRY:
This is retry {attempt} of {max_attempts}.
The previous response did not contain enough valid questions.
Generate a COMPLETE replacement set of exactly {requested_count}
questions.

Do not copy invalid questions from the previous attempt.
Be extremely strict about the requested question type.
For MCQ, use exactly 4 unique options and make the answer exactly
match one option.
For True / False, use only True or False and options exactly
["True", "False"].
For Short Answer, use an empty options array and a non-empty answer.
Return ONLY the JSON object.
"""

        try:
            raw = call_ai(prompt + retry_instruction)
            parsed = parse_ai_json(raw)

            items = (
                parsed.get("questions", [])
                if isinstance(parsed, dict)
                else parsed
            )

            validated = validate_questions(
                items,
                requested_count,
                selected_type,
            )

            last_valid_count = len(validated)

            if len(validated) >= requested_count:
                print(
                    f"{mode} QUIZ SUCCESS: "
                    f"{len(validated)} valid questions "
                    f"on attempt {attempt}."
                )
                return validated[:requested_count]

            last_error = (
                f"Llama returned {len(validated)} valid "
                f"questions out of {requested_count}."
            )

            print(
                f"{mode} QUIZ RETRY {attempt}: "
                f"{last_error}"
            )

        except HTTPException:
            raise

        except Exception as exc:
            last_error = str(exc)
            print(
                f"{mode} QUIZ RETRY {attempt}: "
                f"AI response could not be parsed: {exc!r}"
            )

    raise HTTPException(
        status_code=502,
        detail=(
            f"Llama generated only {last_valid_count} valid "
            f"{selected_type} questions out of {requested_count} "
            f"after {max_attempts} attempts. "
            "Please try again."
        ),
    )


# ============================================================
# RAG CHUNKING
# ============================================================

def chunk_text(
    text: str,
    chunk_size: int = 1200,
    overlap: int = 200,
):
    cleaned = " ".join(
        text.split()
    )

    if not cleaned:
        return []

    chunks = []

    start = 0
    length = len(cleaned)

    while start < length:

        end = min(
            start + chunk_size,
            length,
        )

        chunk = cleaned[
            start:end
        ].strip()

        if chunk:
            chunks.append(chunk)

        if end >= length:
            break

        start = max(
            end - overlap,
            start + 1,
        )

    return chunks


# ============================================================
# EMBEDDINGS
# ============================================================

def embed_texts(
    texts: List[str],
):
    if not texts:
        return []

    data = ollama_post(
        "/api/embed",
        {
            "model": OLLAMA_EMBED_MODEL,
            "input": texts,
            "truncate": True,
        },
        timeout=300,
    )

    embeddings = data.get(
        "embeddings"
    )

    if (
        not embeddings
        or len(embeddings) != len(texts)
    ):
        raise HTTPException(
            status_code=502,
            detail=(
                "Ollama did not return valid embeddings. "
                "Check nomic-embed-text."
            ),
        )

    return embeddings


def cosine_similarity(
    a: np.ndarray,
    b: np.ndarray,
):
    denom = (
        np.linalg.norm(a)
        * np.linalg.norm(b)
    )

    if denom == 0:
        return 0.0

    return float(
        np.dot(a, b) / denom
    )


def build_rag_index(
    source_text: str,
    source_name: str,
):
    content_hash = hashlib.sha256(
        source_text.encode("utf-8")
    ).hexdigest()

    if content_hash in RAG_CACHE:
        return RAG_CACHE[
            content_hash
        ]

    chunks = chunk_text(
        source_text
    )

    if not chunks:
        raise HTTPException(
            status_code=400,
            detail=(
                "No usable text was found "
                "in the document."
            ),
        )

    print(
        f"RAG: creating {len(chunks)} "
        f"chunks for {source_name}"
    )

    embeddings = embed_texts(
        chunks
    )

    index = {
        "source_name": source_name,
        "chunks": chunks,
        "embeddings": np.asarray(
            embeddings,
            dtype=np.float32,
        ),
    }

    RAG_CACHE[
        content_hash
    ] = index

    if len(RAG_CACHE) > 5:

        oldest_key = next(
            iter(RAG_CACHE)
        )

        if oldest_key != content_hash:
            del RAG_CACHE[
                oldest_key
            ]

    return index


def retrieve_context(
    index,
    queries: List[str],
    top_k: int = 10,
):
    chunks = index["chunks"]

    matrix = index["embeddings"]

    query_embeddings = embed_texts(
        queries
    )

    scores = np.zeros(
        len(chunks),
        dtype=np.float32,
    )

    for query_vector in query_embeddings:

        q = np.asarray(
            query_vector,
            dtype=np.float32,
        )

        query_scores = np.array(
            [
                cosine_similarity(
                    q,
                    row,
                )
                for row in matrix
            ],
            dtype=np.float32,
        )

        scores = np.maximum(
            scores,
            query_scores,
        )

    top_k = min(
        top_k,
        len(chunks),
    )

    ranked_indices = np.argsort(
        scores
    )[::-1][:top_k]

    retrieved = []

    for idx in ranked_indices:

        retrieved.append(
            {
                "chunk_id": int(idx),
                "score": float(
                    scores[idx]
                ),
                "text": chunks[idx],
            }
        )

    return retrieved


def build_retrieval_queries(
    req: QuizRequest,
):
    type_query = {
        "MCQ":
            "important concepts, definitions, comparisons, examples and distinctions",

        "True / False":
            "facts, definitions, properties, relationships and statements that can be judged true or false",

        "Short Answer":
            "important definitions, explanations, processes and key concepts",

        "Mixed":
            "important definitions, concepts, examples, processes, comparisons and relationships",
    }[
        req.question_type
    ]

    difficulty_query = {
        "Easy":
            "basic concepts and direct explanations",

        "Medium":
            "conceptual understanding, comparisons and applications",

        "Hard":
            "deeper relationships, reasoning, edge cases and detailed explanations",

        "Mixed":
            "a mixture of basic, conceptual and deeper material",
    }[
        req.difficulty
    ]

    return [
        f"Study material about {type_query}.",

        f"Study material suitable for {difficulty_query} quiz questions.",

        "Most important examinable concepts and definitions.",

        "Processes, examples, relationships, comparisons, formulas, and key facts.",
    ]


def format_retrieved_context(
    retrieved,
):
    parts = []

    for item in retrieved:

        parts.append(
            f"[SOURCE CHUNK {item['chunk_id']} "
            f"| relevance={item['score']:.3f}]\n"
            f"{item['text']}"
        )

    return "\n\n".join(parts)


# ============================================================
# QUIZ GENERATION HELPERS
# ============================================================

def question_type_instruction(
    question_type: str,
):
    return {
        "MCQ":
            "Generate ONLY Multiple Choice Questions.",

        "True / False":
            "Generate ONLY True / False questions.",

        "Short Answer":
            "Generate ONLY Short Answer questions.",

        "Mixed":
            (
                "Generate a balanced mixture of "
                "MCQ, True / False, and Short Answer questions."
            ),
    }[question_type]


def difficulty_instruction(
    difficulty: str,
):
    return {
        "Easy":
            "Test basic concepts, definitions and direct understanding.",

        "Medium":
            "Test conceptual understanding, comparisons, applications and important ideas.",

        "Hard":
            "Test deeper reasoning, relationships, applications, edge cases and detailed understanding.",

        "Mixed":
            "Use a mixture of easy, medium and hard questions.",
    }[difficulty]


def build_topic_prompt(
    req: QuizRequest,
    topic: str,
):
    if req.question_type == "MCQ":

        type_rules = """
Every question MUST be an MCQ.

The "type" MUST be exactly:
"MCQ"

The "options" field MUST contain exactly 4 unique options.

The "answer" MUST exactly match one of the options.

DO NOT generate True / False questions.

DO NOT generate Short Answer questions.
"""

    elif req.question_type == "True / False":

        type_rules = """
Every question MUST be a True / False question.

The "type" MUST be exactly:
"True / False"

The "options" field MUST be exactly:
["True", "False"]

The "answer" MUST be exactly:
"True"
or
"False"

DO NOT generate MCQ questions.

DO NOT generate Short Answer questions.
"""

    elif req.question_type == "Short Answer":

        type_rules = """
Every question MUST be a Short Answer question.

The "type" MUST be exactly:
"Short Answer"

The "options" field MUST be exactly:
[]

The "answer" MUST contain a concise,
correct and NON-EMPTY answer.

DO NOT generate MCQ questions.

DO NOT generate True / False questions.
"""

    else:

        type_rules = """
Generate a balanced mixture of:

- MCQ
- True / False
- Short Answer

MCQ:
- Exactly 4 unique options.
- Answer must match an option.

True / False:
- Options must be ["True", "False"].
- Answer must be exactly "True" or "False".

Short Answer:
- Options must be [].
- Answer must be non-empty.
"""

    return f"""
You are QuizMate AI.

Create EXACTLY {req.count} NEW educational quiz questions.

TOPIC:
{topic}

QUESTION TYPE:
{question_type_instruction(req.question_type)}

DIFFICULTY:
{difficulty_instruction(req.difficulty)}

============================================================
CRITICAL QUESTION TYPE RULE
============================================================

The selected question type is:

"{req.question_type}"

{type_rules}

============================================================
GENERAL RULES
============================================================

1. Questions must be specifically about the requested topic.

2. Do not repeat questions.

3. Make every question meaningfully different.

4. Generate EXACTLY {req.count} questions.

5. Every question MUST contain:

   type
   question
   options
   answer
   explanation
   keywords

6. Never omit any field.

7. Every answer MUST be non-empty.

8. Generate educationally accurate questions.

9. Do not use markdown.

10. Do not return ```json.

11. Return ONLY valid JSON.

12. Generate new questions every time.

============================================================
FINAL VALIDATION BEFORE RESPONSE
============================================================

Before returning the JSON, verify:

- Exactly {req.count} questions exist.
- Every question uses the selected type.
- No question uses a different type.
- Every answer is non-empty.
- All required fields exist.

============================================================
OUTPUT FORMAT
============================================================

Return:

{{
  "questions": [
    {{
      "type": "{req.question_type}",
      "question": "Question text",
      "options": [],
      "answer": "Correct answer",
      "explanation": "Brief explanation.",
      "keywords": []
    }}
  ]
}}
"""


# ============================================================
# QUIZ GENERATION
# ============================================================

def generate_questions(
    req: QuizRequest,
):

    # ========================================================
    # TOPIC-ONLY MODE
    #
    # IMPORTANT:
    # Direct topic mode does NOT search data/topics/*.txt.
    #
    # Example:
    # topic = "iot"
    #
    # It directly sends "iot" to Ollama.
    # ========================================================

    if req.topic and not req.source_text:

        topic = req.topic.strip()

        if len(topic) < 2:
            raise HTTPException(
                status_code=400,
                detail="Please enter a valid topic.",
            )

        prompt = build_topic_prompt(
            req,
            topic,
        )

        return generate_validated_questions(
            prompt=prompt,
            requested_count=req.count,
            selected_type=req.question_type,
            mode="TOPIC",
            max_attempts=3,
        )

    # ========================================================
    # DOCUMENT / RAG MODE
    # ========================================================

    if req.source_text:

        source_text = (
            req.source_text.strip()
        )

        source_name = (
            req.source_name
            or "Uploaded Document"
        )

    elif req.topic:

        # This is kept for EXISTING topic-file functionality.
        # Direct topic mode above returns before reaching here.
        source_name, source_text = (
            get_topic_source(
                req.topic
            )
        )

    else:

        raise HTTPException(
            status_code=400,
            detail=(
                "Please upload a document "
                "or select a topic."
            ),
        )

    if len(source_text) < 100:

        raise HTTPException(
            status_code=400,
            detail=(
                "The selected document/topic "
                "does not contain enough study material."
            ),
        )

    index = build_rag_index(
        source_text,
        source_name,
    )

    retrieval_queries = (
        build_retrieval_queries(req)
    )

    top_k = min(
        max(
            8,
            req.count // 2 + 5,
        ),
        15,
        len(index["chunks"]),
    )

    retrieved = retrieve_context(
        index,
        retrieval_queries,
        top_k=top_k,
    )

    if not retrieved:

        raise HTTPException(
            status_code=502,
            detail=(
                "RAG could not retrieve "
                "relevant content."
            ),
        )

    context = format_retrieved_context(
        retrieved
    )

    # ========================================================
    # DOCUMENT TYPE RULES
    # ========================================================

    if req.question_type == "MCQ":

        document_type_rules = """
Generate ONLY MCQ questions.

Every "type" MUST be exactly "MCQ".

Each question MUST contain exactly 4 unique options.

The answer MUST exactly match one option.

Do NOT generate True / False.

Do NOT generate Short Answer.
"""

    elif req.question_type == "True / False":

        document_type_rules = """
Generate ONLY True / False questions.

Every "type" MUST be exactly "True / False".

Every question MUST have:

"options": ["True", "False"]

The answer MUST be exactly "True" or "False".

Do NOT generate MCQ.

Do NOT generate Short Answer.
"""

    elif req.question_type == "Short Answer":

        document_type_rules = """
Generate ONLY Short Answer questions.

Every "type" MUST be exactly "Short Answer".

Every question MUST have:

"options": []

Every answer MUST contain a concise,
correct and non-empty answer.

Do NOT generate MCQ.

Do NOT generate True / False.
"""

    else:

        document_type_rules = """
Generate a balanced mixture of:

- MCQ
- True / False
- Short Answer

MCQ:
- Exactly 4 unique options.
- Answer must match one option.

True / False:
- Options must be ["True", "False"].
- Answer must be exactly "True" or "False".

Short Answer:
- Options must be [].
- Answer must be non-empty.
"""

    prompt = f"""
You are QuizMate AI.

Generate EXACTLY {req.count} NEW quiz questions.

Use ONLY the retrieved study context.

QUESTION TYPE:
{question_type_instruction(req.question_type)}

DIFFICULTY:
{difficulty_instruction(req.difficulty)}

============================================================
STRICT QUESTION TYPE RULES
============================================================

{document_type_rules}

============================================================
DOCUMENT RULES
============================================================

1. Use ONLY information explicitly supported by the context.

2. Do NOT use outside knowledge.

3. Do NOT invent facts.

4. Do NOT repeat questions.

5. Make every question meaningfully different.

6. Every answer must be supported by the context.

7. Every question MUST contain:

   type
   question
   options
   answer
   explanation
   keywords

8. Never omit any field.

9. Every answer MUST be non-empty.

10. Return ONLY valid JSON.

11. Do not return markdown.

12. Do not return ```json.

============================================================
FINAL VALIDATION
============================================================

Before returning the JSON, verify:

- Exactly {req.count} questions exist.
- Every question uses the selected type.
- No question has a different type.
- Every answer is non-empty.
- All required fields exist.

============================================================
OUTPUT
============================================================

Return:

{{
  "questions": [
    {{
      "type": "{req.question_type}",
      "question": "Question",
      "options": [],
      "answer": "Correct answer",
      "explanation": "Explanation from context.",
      "keywords": []
    }}
  ]
}}

============================================================
RETRIEVED STUDY CONTEXT
============================================================

<RETRIEVED_CONTEXT>
{context}
</RETRIEVED_CONTEXT>
"""

    return generate_validated_questions(
        prompt=prompt,
        requested_count=req.count,
        selected_type=req.question_type,
        mode="DOCUMENT",
        max_attempts=3,
    )


# ============================================================
# QUIZ SCORING
# ============================================================

def score_questions(
    questions,
    answers,
):
    score = 0

    for i, question in enumerate(
        questions
    ):

        user_answer = str(
            answers.get(
                str(i),
                "",
            )
        ).strip().lower()

        correct_answer = (
            str(
                question.answer
            )
            .strip()
            .lower()
        )

        # --------------------------------------------
        # No answer
        # --------------------------------------------

        if not user_answer:
            continue

        # --------------------------------------------
        # MCQ / True False
        # --------------------------------------------

        if question.type in {
            "MCQ",
            "True / False",
        }:

            if user_answer == correct_answer:
                score += 1

            continue

        # --------------------------------------------
        # Short Answer
        # --------------------------------------------

        if question.type == "Short Answer":

            keywords = [
                str(k).strip().lower()
                for k in question.keywords
                if str(k).strip()
            ]

            if keywords:

                matched = sum(
                    1
                    for keyword in keywords
                    if keyword in user_answer
                )

                if matched >= 1:
                    score += 1

            else:

                answer_words = [
                    word
                    for word in correct_answer.split()
                    if len(word) > 3
                ]

                if answer_words:

                    matched = sum(
                        1
                        for word in answer_words
                        if word in user_answer
                    )

                    if (
                        matched
                        >= max(
                            1,
                            len(answer_words) // 3,
                        )
                    ):
                        score += 1

                elif (
                    user_answer
                    == correct_answer
                ):
                    score += 1

    return score


# ============================================================
# HEALTH
# ============================================================

@app.get("/")
def health():
    return {
        "name": "QuizMate AI",
        "status": "ok",
        "ollama": OLLAMA_BASE_URL,
        "model": OLLAMA_MODEL,
    }


# ============================================================
# TOPIC API
# ============================================================

@app.get("/api/topics")
def topics():

    return [
        {
            "name": name,
            "description": (
                "Practice important concepts with AI."
            ),
        }
        for name in load_topics()
    ]


@app.get("/api/topics/{topic_name}")
def topic_content(
    topic_name: str,
):

    data = load_topics()

    if topic_name not in data:

        raise HTTPException(
            status_code=404,
            detail="Topic not found",
        )

    return {
        "name": topic_name,
        "text": data[topic_name],
    }


# ============================================================
# AUTH REGISTER
# ============================================================

@app.post(
    "/api/auth/register",
    response_model=TokenOut,
)
def register(
    req: AuthIn,
):

    if (
        not req.name
        or len(req.password) < 6
    ):

        raise HTTPException(
            status_code=400,
            detail=(
                "Name and password "
                "(6+ characters) are required."
            ),
        )

    email = req.email.strip().lower()

    con = get_db()

    try:

        cur = con.execute(
            """
            INSERT INTO users(
                name,
                email,
                password_hash
            )
            VALUES(?,?,?)
            """,
            (
                req.name.strip(),
                email,
                hash_password(
                    req.password
                ),
            ),
        )

        con.commit()

        user_id = cur.lastrowid

    except sqlite3.IntegrityError:

        con.close()

        raise HTTPException(
            status_code=409,
            detail=(
                "An account with this "
                "email already exists."
            ),
        )

    finally:

        try:
            con.close()
        except Exception:
            pass

    return {
        "token": make_token(
            user_id,
            email,
        ),
        "user": {
            "id": user_id,
            "name": req.name.strip(),
            "email": email,
        },
    }


# ============================================================
# AUTH LOGIN
# ============================================================

@app.post(
    "/api/auth/login",
    response_model=TokenOut,
)
def login(
    req: AuthIn,
):

    email = req.email.strip().lower()

    con = get_db()

    row = con.execute(
        """
        SELECT *
        FROM users
        WHERE email=?
        """,
        (email,),
    ).fetchone()

    con.close()

    if (
        not row
        or not verify_password(
            req.password,
            row["password_hash"],
        )
    ):

        raise HTTPException(
            status_code=401,
            detail="Invalid email or password.",
        )

    return {
        "token": make_token(
            row["id"],
            row["email"],
        ),
        "user": {
            "id": row["id"],
            "name": row["name"],
            "email": row["email"],
        },
    }


# ============================================================
# GENERATE QUIZ
# ============================================================

@app.post(
    "/api/quiz/generate",
    response_model=QuizOut,
)
def quiz_generate(
    req: QuizRequest,
    user=Depends(current_user),
):

    questions = generate_questions(
        req
    )

    return {
        "questions": questions
    }


# ============================================================
# SUBMIT QUIZ
# ============================================================

@app.post("/api/quiz/submit")
def quiz_submit(
    req: SubmitRequest,
    user=Depends(current_user),
):

    try:

        if not req.questions:
            raise HTTPException(
                status_code=400,
                detail="No quiz questions supplied.",
            )

        if not isinstance(
            req.answers,
            dict,
        ):
            raise HTTPException(
                status_code=400,
                detail="Answers must be an object.",
            )

        print(
            "\n========== QUIZ SUBMIT =========="
        )

        print(
            "USER ID:",
            user["id"],
        )

        print(
            "SOURCE:",
            req.source_name,
        )

        print(
            "QUESTION COUNT:",
            len(req.questions),
        )

        print(
            "ANSWER COUNT:",
            len(req.answers),
        )

        print(
            "ANSWERS:",
            req.answers,
        )

        score = score_questions(
            req.questions,
            req.answers,
        )

        total = len(
            req.questions
        )

        con = get_db()

        con.execute(
            """
            INSERT INTO results(
                user_id,
                topic,
                score,
                total,
                difficulty,
                created_at
            )
            VALUES(?,?,?,?,?,?)
            """,
            (
                user["id"],
                req.source_name,
                score,
                total,
                req.difficulty,
                datetime.now(
                    timezone.utc
                ).isoformat(),
            ),
        )

        con.commit()
        con.close()

        percentage = (
            round(
                score / total * 100
            )
            if total
            else 0
        )

        print(
            "SCORE:",
            score,
            "/",
            total,
        )

        print(
            "PERCENTAGE:",
            percentage,
        )

        print(
            "=================================\n"
        )

        return {
            "score": score,
            "total": total,
            "percentage": percentage,
        }

    except HTTPException:
        raise

    except Exception as e:

        print(
            "\n========== QUIZ SUBMIT ERROR =========="
        )

        print(
            "ERROR:",
            repr(e),
        )

        import traceback

        traceback.print_exc()

        print(
            "========================================\n"
        )

        raise HTTPException(
            status_code=500,
            detail=(
                f"Quiz submission failed: {str(e)}"
            ),
        )


# ============================================================
# RESULTS
# ============================================================

@app.get("/api/results")
def results(
    user=Depends(current_user),
):

    con = get_db()

    rows = [
        dict(row)
        for row in con.execute(
            """
            SELECT
                topic,
                score,
                total,
                difficulty,
                created_at
            FROM results
            WHERE user_id=?
            ORDER BY id DESC
            """,
            (user["id"],),
        ).fetchall()
    ]

    con.close()

    return rows


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_text_from_file(
    file_path: str,
    filename: str,
) -> str:

    extension = (
        filename
        .lower()
        .split(".")[-1]
    )

    # ========================================================
    # PDF
    # ========================================================

    if extension == "pdf":

        text_parts = []

        doc = fitz.open(
            file_path
        )

        try:

            for page in doc:

                text = page.get_text()

                if text:
                    text_parts.append(
                        text
                    )

        finally:

            doc.close()

        return "\n".join(
            text_parts
        ).strip()

    # ========================================================
    # DOCX
    # ========================================================

    if extension == "docx":

        document = Document(
            file_path
        )

        text_parts = []

        for paragraph in (
            document.paragraphs
        ):

            text = (
                paragraph.text
                .strip()
            )

            if text:
                text_parts.append(
                    text
                )

        for table in (
            document.tables
        ):

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = (
                        cell.text
                        .strip()
                    )

                    if cell_text:
                        row_text.append(
                            cell_text
                        )

                if row_text:
                    text_parts.append(
                        " | ".join(
                            row_text
                        )
                    )

        return "\n".join(
            text_parts
        ).strip()

    raise ValueError(
        "Unsupported file type. "
        "Only PDF and DOCX files are supported."
    )


@app.post("/api/pdf/extract")
async def extract_document(
    file: UploadFile = File(...),
):

    filename = file.filename or ""

    extension = (
        filename
        .lower()
        .split(".")[-1]
    )

    if extension not in {
        "pdf",
        "docx",
    }:

        raise HTTPException(
            status_code=400,
            detail=(
                "Only PDF and DOCX "
                "files are supported."
            ),
        )

    temp_path = None

    try:

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=f".{extension}",
        ) as temp_file:

            content = await file.read()

            if not content:
                raise HTTPException(
                    status_code=400,
                    detail="Uploaded file is empty.",
                )

            # Keep the existing UI's 10 MB upload limit enforced
            # on the backend as well.
            max_file_size = 10 * 1024 * 1024

            if len(content) > max_file_size:
                raise HTTPException(
                    status_code=413,
                    detail="File is too large. Maximum size is 10 MB.",
                )

            temp_file.write(
                content
            )

            temp_path = (
                temp_file.name
            )

        text = extract_text_from_file(
            temp_path,
            filename,
        )

        if not text.strip():

            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not extract any text "
                    "from this document."
                ),
            )

        print(
            "DOCUMENT EXTRACTED:",
            filename,
            "characters:",
            len(text),
        )

        return {
            "filename": filename,
            "text": text,
            "file_type": extension,
        }

    except HTTPException:
        raise

    except Exception as e:

        print(
            "Document extraction error:",
            repr(e),
        )

        raise HTTPException(
            status_code=500,
            detail=(
                f"Failed to process document: {str(e)}"
            ),
        )

    finally:

        if (
            temp_path
            and os.path.exists(temp_path)
        ):

            try:
                os.remove(
                    temp_path
                )
            except Exception:
                pass


# ============================================================
# SUMMARY
# ============================================================

@app.post("/api/summary")
def summary(
    req: SummaryRequest,
    user=Depends(current_user),
):

    text = req.text.strip()

    if len(text) < 20:

        raise HTTPException(
            status_code=400,
            detail=(
                "Document text is too short "
                "to summarize."
            ),
        )

    prompt = f"""
You are QuizMate AI.

Summarize ONLY the document below.

SUMMARY STYLE:
{req.mode}

IMPORTANT:

1. Use only information from the document.
2. Do not invent facts.
3. The summary must not be empty.
4. Make the summary useful for studying.
5. Return ONLY valid JSON.
6. Do not return markdown.
7. Do not return ```json.

Return EXACTLY:

{{
  "summary": "Your complete summary here."
}}

DOCUMENT:

<DOCUMENT>
{text[:40000]}
</DOCUMENT>
"""

    try:

        raw = call_ai(
            prompt
        )

        parsed = parse_ai_json(
            raw
        )

        if not isinstance(
            parsed,
            dict,
        ):

            raise ValueError(
                "Summary response is not a JSON object."
            )

        summary_text = str(
            parsed.get(
                "summary",
                "",
            )
        ).strip()

        # ====================================================
        # IMPORTANT FALLBACK
        # ====================================================

        if not summary_text:

            print(
                "SUMMARY JSON WAS EMPTY. "
                "Retrying without JSON mode."
            )

            fallback_payload = {
                "model": OLLAMA_MODEL,

                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "You are a document summarization "
                            "assistant. Summarize only the supplied "
                            "document. Do not invent information."
                        ),
                    },
                    {
                        "role": "user",
                        "content": f"""
Summarize this document.

Style:
{req.mode}

Document:
{text[:40000]}
""",
                    },
                ],

                "stream": False,

                "options": {
                    "temperature": 0.2,
                },
            }

            fallback_data = ollama_post(
                "/api/chat",
                fallback_payload,
                timeout=300,
            )

            summary_text = (
                fallback_data
                .get("message", {})
                .get("content", "")
                .strip()
            )

        if not summary_text:

            raise HTTPException(
                status_code=502,
                detail=(
                    "Ollama returned an empty summary."
                ),
            )

        print(
            "SUMMARY SUCCESS:",
            len(summary_text),
            "characters",
        )

        return {
            "summary": summary_text
        }

    except HTTPException:
        raise

    except Exception as e:

        print(
            "SUMMARY ERROR:",
            repr(e),
        )

        import traceback

        traceback.print_exc()

        raise HTTPException(
            status_code=502,
            detail=(
                "Failed to generate document summary."
            ),
        )
