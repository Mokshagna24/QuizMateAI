import os
import tempfile

import fitz
from docx import Document
from fastapi import File, HTTPException, UploadFile

from ..core.app import app


def extract_text_from_file(
    file_path: str,
    filename: str,
) -> str:

    extension = (
        filename
        .lower()
        .split(".")[-1]
    )

    # ========================================================
    # PDF
    # ========================================================

    if extension == "pdf":

        text_parts = []

        doc = fitz.open(
            file_path
        )

        try:

            for page in doc:

                text = page.get_text()

                if text:
                    text_parts.append(
                        text
                    )

        finally:

            doc.close()

        return "\n".join(
            text_parts
        ).strip()

    # ========================================================
    # DOCX
    # ========================================================

    if extension == "docx":

        document = Document(
            file_path
        )

        text_parts = []

        for paragraph in (
            document.paragraphs
        ):

            text = (
                paragraph.text
                .strip()
            )

            if text:
                text_parts.append(
                    text
                )

        for table in (
            document.tables
        ):

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = (
                        cell.text
                        .strip()
                    )

                    if cell_text:
                        row_text.append(
                            cell_text
                        )

                if row_text:
                    text_parts.append(
                        " | ".join(
                            row_text
                        )
                    )

        return "\n".join(
            text_parts
        ).strip()

    raise ValueError(
        "Unsupported file type. "
        "Only PDF and DOCX files are supported."
    )


@app.post("/api/pdf/extract")
async def extract_document(
    file: UploadFile = File(...),
):

    filename = file.filename or ""

    extension = (
        filename
        .lower()
        .split(".")[-1]
    )

    if extension not in {
        "pdf",
        "docx",
    }:

        raise HTTPException(
            status_code=400,
            detail=(
                "Only PDF and DOCX "
                "files are supported."
            ),
        )

    temp_path = None

    try:

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=f".{extension}",
        ) as temp_file:

            content = await file.read()

            if not content:
                raise HTTPException(
                    status_code=400,
                    detail="Uploaded file is empty.",
                )

            # Keep the existing UI's 10 MB upload limit enforced
            # on the backend as well.
            max_file_size = 10 * 1024 * 1024

            if len(content) > max_file_size:
                raise HTTPException(
                    status_code=413,
                    detail="File is too large. Maximum size is 10 MB.",
                )

            temp_file.write(
                content
            )

            temp_path = (
                temp_file.name
            )

        text = extract_text_from_file(
            temp_path,
            filename,
        )

        if not text.strip():

            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not extract any text "
                    "from this document."
                ),
            )

        print(
            "DOCUMENT EXTRACTED:",
            filename,
            "characters:",
            len(text),
        )

        return {
            "filename": filename,
            "text": text,
            "file_type": extension,
        }

    except HTTPException:
        raise

    except Exception as e:

        print(
            "Document extraction error:",
            repr(e),
        )

        raise HTTPException(
            status_code=500,
            detail=(
                f"Failed to process document: {str(e)}"
            ),
        )

    finally:

        if (
            temp_path
            and os.path.exists(temp_path)
        ):

            try:
                os.remove(
                    temp_path
                )
            except Exception:
                pass
